"""导出 —— Markdown / Docx（补充计划第 6 节）。

汇总任务的清洗稿、结构拆解、三类改写，写入 data/exports 并登记 Artifact。
"""

from pathlib import Path

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.paths import exports_dir
from app.models import Artifact, ArtifactKind, Task

# 导出中展示的产物顺序与标题
_SECTIONS = [
    (ArtifactKind.REVISED_TRANSCRIPT, "转写稿（人工修订）"),
    (ArtifactKind.RAW_TRANSCRIPT, "转写稿（原始）"),
    (ArtifactKind.CLEANED, "清洗稿"),
    (ArtifactKind.BREAKDOWN, "结构拆解"),
    (ArtifactKind.REWRITE_ORAL, "改写 · 口播稿"),
    (ArtifactKind.REWRITE_NOTE, "改写 · 种草笔记"),
    (ArtifactKind.REWRITE_TITLE, "改写 · 标题与开头钩子"),
]


def _collect(session: Session, task_id: str) -> dict[ArtifactKind, str]:
    arts = session.scalars(select(Artifact).where(Artifact.task_id == task_id)).all()
    return {a.kind: a.text for a in arts if a.text}


def _upsert_export(
    session: Session, task_id: str, kind: ArtifactKind, path: str
) -> None:
    art = session.scalar(
        select(Artifact).where(Artifact.task_id == task_id, Artifact.kind == kind)
    )
    if art is None:
        session.add(Artifact(task_id=task_id, kind=kind, file_path=path))
    else:
        art.file_path = path
    session.flush()


def build_markdown(session: Session, task: Task) -> str:
    texts = _collect(session, task.id)
    lines = [f"# 短视频文案 · {task.video_title or task.id}", ""]
    if task.source_url:
        lines.append(f"来源：{task.source_url}")
    if task.author:
        lines.append(f"作者：{task.author}")
    lines.append("")
    # 修订稿存在时不再重复输出原始稿
    skip_raw = ArtifactKind.REVISED_TRANSCRIPT in texts
    for kind, title in _SECTIONS:
        if kind == ArtifactKind.RAW_TRANSCRIPT and skip_raw:
            continue
        if kind in texts:
            lines.append(f"## {title}")
            lines.append("")
            lines.append(texts[kind])
            lines.append("")
    return "\n".join(lines)


def export_markdown(session: Session, task: Task) -> str:
    content = build_markdown(session, task)
    out = exports_dir() / f"{task.id}.md"
    out.write_text(content, encoding="utf-8")
    _upsert_export(session, task.id, ArtifactKind.EXPORT_MD, str(out))
    return str(out)


def export_docx(session: Session, task: Task) -> str:
    from docx import Document

    texts = _collect(session, task.id)
    doc = Document()
    doc.add_heading(f"短视频文案 · {task.video_title or task.id}", level=0)
    if task.source_url:
        doc.add_paragraph(f"来源：{task.source_url}")
    if task.author:
        doc.add_paragraph(f"作者：{task.author}")

    skip_raw = ArtifactKind.REVISED_TRANSCRIPT in texts
    for kind, title in _SECTIONS:
        if kind == ArtifactKind.RAW_TRANSCRIPT and skip_raw:
            continue
        if kind in texts:
            doc.add_heading(title, level=1)
            for para in texts[kind].split("\n\n"):
                doc.add_paragraph(para)

    out: Path = exports_dir() / f"{task.id}.docx"
    doc.save(str(out))
    _upsert_export(session, task.id, ArtifactKind.EXPORT_DOCX, str(out))
    return str(out)
